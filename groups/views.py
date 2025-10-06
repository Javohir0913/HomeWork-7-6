import os
import zipfile
from pathlib import Path
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.contrib.auth.mixins import LoginRequiredMixin, UserPassesTestMixin
from django.shortcuts import render, redirect
from django.urls import reverse_lazy
from django.views import View
from django.views.generic import ListView, CreateView, UpdateView
import pandas as pd
from django.http import HttpResponse
from datetime import timedelta
from django.utils import timezone
import pytz
from fpdf import FPDF
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches
from .helper import data, data_2
from django.conf import settings

from .forms import GroupViewForm
from .models import Student, Attendance, Group, Honadon, GroupView


# Create your views here.
class GroupList(LoginRequiredMixin, ListView):
    model = GroupView
    template_name = 'group/my_group.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        groups = GroupView.objects.filter(user=self.request.user.id)

        # Create a list to store group information and student counts
        all_info = []

        for group_view in groups:
            for group in group_view.groups_list.all():
                count_people = Student.objects.filter(stundet_group=group).count()  # O'zgarish
                all_info.append({'group': group, 'count_people': count_people})

        context['all_info'] = all_info
        return context


class GroupAttendanceView(LoginRequiredMixin, View):
    template_name = 'group/my_group_detail.html'

    def get(self, request, *args, **kwargs):
        today_start = timezone.now().replace(hour=0, minute=0, second=0, microsecond=0)
        today_end = today_start + timedelta(days=1)
        group_id = kwargs.get('pk')
        group = Group.objects.get(id=group_id)
        students = Student.objects.filter(stundet_group=group_id)
        zxc = Attendance.objects.filter(data_day__range=(today_start, today_end), group_id=group).first()
        if zxc is None:
            context = {
                'students': students,
                'group': group,  # Group ob'ektini shablonga yuboryapmiz
                'para1': 'True',
                'para2': 'False',
                'para3': 'False'
            }
        elif len(zxc.para2) == 0:
            context = {
                'students': students,
                'group': group,  # Group ob'ektini shablonga yuboryapmiz
                'para1': 'False',
                'para2': 'True',
                'para3': 'False'
            }
        elif len(zxc.para3) == 0:
            context = {
                'students': students,
                'group': group,  # Group ob'ektini shablonga yuboryapmiz
                'para1': 'False',
                'para2': 'False',
                'para3': 'True'
            }
        else:
            context = {
                'students': students,
                'group': group,  # Group ob'ektini shablonga yuboryapmiz
                'para1': 'False',
                'para2': 'False',
                'para3': 'False'
            }
        return render(request, self.template_name, context)

    def post(self, request, *args, **kwargs):
        today_start = timezone.now().replace(hour=0, minute=0, second=0, microsecond=0)
        today_end = today_start + timedelta(days=1)
        group_id = kwargs.get('pk')
        group = Group.objects.get(id=group_id)
        students = Student.objects.filter(stundet_group=group_id)

        # Form ma'lumotlarini bu yerda qayta ishlash
        for student in students:
            para1 = request.POST.get(f'para1_{student.id}')
            para2 = request.POST.get(f'para2_{student.id}')
            para3 = request.POST.get(f'para3_{student.id}')
            # Eski ma'lumotlarni olish (agar mavjud bo'lsa)
            attendance, created = Attendance.objects.get_or_create(
                student=student, group_id=group, data_day__range=(today_start, today_end),
            )
            # Eski qiymatlarni saqlab, faqat kelgan yangilarini yangilash
            if para1:
                attendance.para1 = para1
            if para2:
                attendance.para2 = para2
            if para3:
                attendance.para3 = para3

            # O'zgarishlarni saqlash
            attendance.save()

        # Form ma'lumotlarini qayta ishlagandan keyin sahifani qayta yuklash yoki boshqa sahifaga o'tish
        return redirect('group_list')  # To'g'ri URL manzilini qo'yishingiz kerak


@login_required
def export_attendance_to_excel(request, pk):
    # Bugungi kunning boshlanish va tugash vaqtini olish
    today_start = timezone.now().replace(hour=0, minute=0, second=0, microsecond=0)
    today_end = today_start + timedelta(days=1)

    # Bugungi sanaga mos ma'lumotlarni olish
    queryset = Attendance.objects.filter(data_day__range=(today_start, today_end), group_id=pk)

    # Mahalliy vaqt zonasini olish
    local_tz = pytz.timezone('Asia/Tashkent')

    # Ma'lumotlarni olish va ism-familiyani birlashtirish
    data = list(
        queryset.values('student__student_name', 'student__student_last_name', 'para1', 'para2', 'para3', 'data_day',
                        'group_id__group_name'))

    # Ism va familiyani birlashtirib yangi ro'yxat yaratish
    for item in data:
        # UTC vaqtni mahalliy vaqtga o'zgartirish
        utc_time = item['data_day'].replace(tzinfo=pytz.utc)
        local_time = utc_time.astimezone(local_tz)
        item['data_day'] = local_time.strftime('%Y-%m-%d %H:%M:%S')  # Vaqt formatlash

        item['full_name'] = item['student__student_last_name'] + ' ' + item['student__student_name']
        del item['student__student_name']
        del item['student__student_last_name']

    if not data:
        messages.warning(request, "Bugungi sana uchun bu guruhda qatnash ma'lumotlari topilmadi.")
        return redirect('group_list')  # Foydalanuvchini kerakli sahifaga yo'naltirish

    # Pandas DataFrame yaratish
    df = pd.DataFrame(data)

    # Ustunlarni tartibga solish (agar kerak bo'lsa)
    df = df[['full_name', 'para1', 'para2', 'para3', 'data_day', 'group_id__group_name']]

    # Excel fayl yaratish
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="{Group.objects.get(id=pk)} {today_start.date()}.xlsx"'

    with pd.ExcelWriter(response, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Attendance')

    return response


def load_students_from_excel(file_path):
    # Excel faylini o'qish (ustunlarsiz, ya'ni header=None)
    df = pd.read_excel(file_path, header=None)
    # Har bir qatorni ma'lumotlar bazasiga saqlash
    for index, row in df.iterrows():
        # Guruhni topish yoki yaratish (D ustun - 3-indeks)
        group, created = Group.objects.get_or_create(group_name=row[3])
        # Talabani yaratish yoki saqlash (A, B, C ustunlar - 0, 1, 2 indekslar)
        student, created = Student.objects.get_or_create(
            student_last_name=row[0],  # A ustun - Familiya
            student_name=row[1],       # B ustun - Ism
            student_father_name=row[2],# C ustun - Otasining ismi
            stundet_group=group
        )

        print(f"Yaratildi: {student}")

    return HttpResponse("Ma'lumotlar yuklandi!")


def excel_to_word(excel_file, output_dir):
    zxc = Path(settings.MEDIA_ROOT)
    output_dir.mkdir(exist_ok=True)  # Papkani yaratish (agar mavjud bo'lmasa)

    df = pd.read_excel(excel_file)
    word_files = []  # Yaratilgan Word fayllarni saqlash uchun ro'yxat

    for index, row in df.iterrows():
        if pd.notna(row[4]) and str(row[4]).strip():
            doc = Document()

            # Sarlavhani qo'shish va formatlash
            title = "Toshkent amaliy fanlar universiteti talabasining ijtimoiy-psixologik pasporti"
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(title)
            run.bold = True
            run.italic = True
            run.font.name = 'Times New Roman'  # Shriftni Times New Roman ga o'rnatish
            run.font.size = Pt(18)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # O'rtaga joylashtirish

            # Yangi qator qo'shish uchun bo'sh paragraph yaratish
            doc.add_paragraph()

            # "1. Talabaning F.I.Sh." matnini qo'shish
            talaba_paragraph = doc.add_paragraph()
            talaba_run = talaba_paragraph.add_run(f"1. Talabaning F.I.Sh.  ")
            talaba_run.font.name = 'Times New Roman'
            talaba_run.font.size = Pt(14)
            talaba_run.bold = True

            # Talaba ismini qo'shish, underline bilan formatlash
            talaba_name_run = talaba_paragraph.add_run(f"{row[4]}")
            talaba_name_run.font.name = 'Times New Roman'
            talaba_name_run.font.size = Pt(14)
            talaba_name_run.underline = True

            # Fakulteti, guruhi, kursi
            talaba_group = doc.add_paragraph()
            talaba_group_run = talaba_group.add_run("2. Fakulteti, Guruhi, Kursi:  ")
            talaba_group_run.font.name = 'Times New Roman'
            talaba_group_run.font.size = Pt(14)
            talaba_group_run.bold = True

            talaba_group_add = talaba_group.add_run(f"{row[125]}, {row[126]} {row[3]}")
            talaba_group_add.font.name = 'Times New Roman'
            talaba_group_add.font.size = Pt(14)
            talaba_group_add.underline = True

            talba_millat = doc.add_paragraph()
            talba_millat_run = talba_millat.add_run("3. Millati:  ")
            talba_millat_run.font.name = 'Times New Roman'
            talba_millat_run.font.size = Pt(14)
            talba_millat_run.font.bold = True

            talba_millat_add = talba_millat.add_run(f"{row[5]}")
            talba_millat_add.font.name = 'Times New Roman'
            talba_millat_add.font.size = Pt(14)
            talba_millat_add.underline = True

            talaba_tg_kun = talba_millat.add_run(" 4. Tug‘ilgan kuni: ")
            talaba_tg_kun.font.name = 'Times New Roman'
            talaba_tg_kun.font.size = Pt(14)
            talaba_tg_kun.font.bold = True

            talaba_tg_kun_add = talba_millat.add_run(f"{row[6]}")
            talaba_tg_kun_add.font.name = 'Times New Roman'
            talaba_tg_kun_add.font.size = Pt(14)
            talaba_tg_kun_add.underline = True

            talaba_manzil = doc.add_paragraph()
            talaba_manzil_run = talaba_manzil.add_run("5. Doimiy yashash manzili: ")
            talaba_manzil_run.font.name = 'Times New Roman'
            talaba_manzil_run.font.size = Pt(14)
            talaba_manzil_run.font.bold = True

            talaba_manzil_add = talaba_manzil.add_run(f"{row[127]}")
            talaba_manzil_add.font.name = 'Times New Roman'
            talaba_manzil_add.font.size = Pt(14)
            talaba_manzil_add.underline = True

            talaba_o_a = doc.add_paragraph()
            talaba_o_a_run = talaba_o_a.add_run("6. Oilaviy ahvoli (oila a’zolari haqida ma’lumot): ")
            talaba_o_a_run.font.name = 'Times New Roman'
            talaba_o_a_run.font.size = Pt(14)
            talaba_o_a_run.font.bold = True

            talaba_o_a_run_add = talaba_o_a.add_run(f"{row[7]}")
            talaba_o_a_run_add.font.name = 'Times New Roman'
            talaba_o_a_run_add.font.size = Pt(14)
            talaba_o_a_run_add.underline = True

            talaba_salomatligi = doc.add_paragraph()
            talaba_salomatligi_run = talaba_salomatligi.add_run("7. Salomatligi to‘g‘risida ma’lumotlar: ")
            talaba_salomatligi_run.font.name = 'Times New Roman'
            talaba_salomatligi_run.font.size = Pt(14)
            talaba_salomatligi_run.font.bold = True

            talaba_salomatligi_add = talaba_salomatligi.add_run(f"{row[8]}")
            talaba_salomatligi_add.font.name = 'Times New Roman'
            talaba_salomatligi_add.font.size = Pt(14)
            talaba_salomatligi_add.underline = True

            talaba_figura_tanla = doc.add_paragraph()
            talaba_figura_tanla_run = talaba_figura_tanla.add_run(
                "8. Yoqtirgan figuralarning ketma-ketligini tanlang (1 dan 5 gacha raqamlab chiqing)")
            talaba_figura_tanla_run.font.name = 'Times New Roman'
            talaba_figura_tanla_run.font.size = Pt(14)
            talaba_figura_tanla_run.font.bold = True
            doc.add_picture(f'{zxc}/image_excel_to_wrod/{row[9][-1]}-figura.png', width=Inches(6))

            talaba_qadrli = doc.add_paragraph()
            talaba_qadrli_run = talaba_qadrli.add_run(
                '9. “Siz uchun hayotda nima qadrli?”, degan savollariga o‘z munosabatingizni bildiring.')
            talaba_qadrli_run.font.name = 'Times New Roman'
            talaba_qadrli_run.font.size = Pt(13)
            talaba_qadrli_run.font.bold = True

            talaba_qadrli_add = doc.add_paragraph()
            talaba_qadrli_add_run = talaba_qadrli_add.add_run(f"{row[10]}")
            talaba_qadrli_add_run.font.name = 'Times New Roman'
            talaba_qadrli_add_run.font.size = Pt(14)
            talaba_qadrli_add_run.underline = True

            talaba_fikri = doc.add_paragraph()
            talaba_fikri_run = talaba_fikri.add_run('Shaxsiy fikringiz: ')
            talaba_fikri_run.font.name = 'Times New Roman'
            talaba_fikri_run.font.size = Pt(14)
            talaba_fikri_run.font.bold = True

            talaba_fikri_add = talaba_fikri.add_run(f"{row[11]}")
            talaba_fikri_add.font.name = 'Times New Roman'
            talaba_fikri_add.font.size = Pt(14)
            talaba_fikri_add.underline = True

            talaba_hayot_mano = doc.add_paragraph()
            talaba_hayot_mano_run = talaba_hayot_mano.add_run(
                'Hayotning ma’nosi deb.............tushunaman savoliga javoblarning birini tanlang.')
            talaba_hayot_mano_run.font.name = 'Times New Roman'
            talaba_hayot_mano_run.font.size = Pt(14)
            talaba_hayot_mano_run.font.bold = True

            talaba_hayot_mano_add = doc.add_paragraph()
            talaba_hayot_mano_add_run = talaba_hayot_mano_add.add_run(f'{row[12]}')
            talaba_hayot_mano_add_run.font.name = 'Times New Roman'
            talaba_hayot_mano_add_run.font.size = Pt(14)
            talaba_hayot_mano_add_run.underline = True

            talaba_hayot_fikri = doc.add_paragraph()
            talaba_hayot_fikri_run = talaba_hayot_fikri.add_run('Shaxsiy fikringiz: ')
            talaba_hayot_fikri_run.font.name = 'Times New Roman'
            talaba_hayot_fikri_run.font.size = Pt(14)
            talaba_hayot_fikri_run.font.bold = True

            talaba_hayot_fikri_add = talaba_hayot_fikri.add_run(f"{row[13]}")
            talaba_hayot_fikri_add.font.name = 'Times New Roman'
            talaba_hayot_fikri_add.font.size = Pt(14)
            talaba_hayot_fikri_add.underline = True

            talaba_tablitsa = doc.add_paragraph()
            talaba_tablitsa_run = talaba_tablitsa.add_run(
                'BERILGAN FIKRGA QO‘SHILSANGIZ, SIZGA TAALLUQLI BO‘LGANLARIGA “1” BALL, QO‘SHILMASANGIZ “0” BALL  QO‘YING.')
            talaba_tablitsa.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            talaba_tablitsa_run.font.bold = True
            talaba_tablitsa_run.font.name = 'Times New Roman'
            talaba_tablitsa_run.font.size = Pt(11)

            table = doc.add_table(rows=len(data), cols=4)
            # Jadvalga matn qo'shish
            for i, row_data in enumerate(data):
                row_table = table.rows[i]

                for j, text in enumerate(row_data):
                    cell = row_table.cells[j]
                    paragraph = cell.paragraphs[0]
                    if text:
                        if str(text)[0].isdigit():
                            run = paragraph.add_run(text)
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(11)
                            run.font.bold = True
                        else:
                            run = paragraph.add_run(text)
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(11)
                        in_else = text
                    else:
                        if in_else in df.columns:
                            column_index = df.columns.get_loc(in_else)
                            if str(row[column_index]):
                                run = paragraph.add_run(str(row[column_index])[0])
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(11)
            doc.add_paragraph()
            table_2_text = doc.add_paragraph()
            table_2_text_run = table_2_text.add_run(
                '  Sizning e’tiboringizga bir qator oddiy savollar havola etilgan. Siz ularga “ha”, “yo‘q”, “ba’zida” deb javob berishingiz mumkin. Savollarga tez va ko‘p o‘ylanmay javob berishingiz so‘raladi.')
            table_2_text_run.font.name = 'Times New Roman'
            table_2_text_run.font.size = Pt(11)
            table_2_text.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            table_2 = doc.add_table(rows=len(data_2), cols=2)
            for i, row_data in enumerate(data_2):
                row_table = table_2.rows[i]  # table_2 ni ishlatishimiz kerak
                for j, text in enumerate(row_data):
                    cell = row_table.cells[j]
                    paragraph = cell.paragraphs[0]
                    if text:
                        run = paragraph.add_run(text)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)
                        in_else = text
                    else:
                        if in_else in df.columns:
                            column_index = df.columns.get_loc(in_else)
                            if str(row[column_index]):
                                run = paragraph.add_run(str(row[column_index]))
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(11)

            # Word hujjatini saqlash
            word_file_path = output_dir / f"{row[4]}.docx"
            doc.save(word_file_path)
            word_files.append(word_file_path)  # Fayl yo'lini ro'yxatga qo'shish

    return word_files  # Yaratilgan Word fayllarni qaytarish


def zip_directory(files, zip_name):
    zip_path = f"{zip_name}.zip"
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for file_path in files:
            zipf.write(file_path, file_path.name)  # Faylni faqat nomi bilan qo'shamiz
    return zip_path


@login_required
def student_excel_word(request):
    if request.method == 'POST' and request.FILES.get('excel_file'):
        excel_file = request.FILES['excel_file']
        output_dir = Path(settings.MEDIA_ROOT) / "word_to_excel"

        # Excel fayldan Word fayllarni yaratish va ularni ro'yxatga olish
        word_files = excel_to_word(excel_file, output_dir)

        # Zip faylni yaratish (faqat yaratgan Word fayllar bilan)
        zip_file_path = zip_directory(word_files, output_dir / "word_documents")

        # Zip faylni foydalanuvchiga qaytarish
        with open(zip_file_path, 'rb') as zip_file:
            response = HttpResponse(zip_file.read(), content_type="application/zip")
            response['Content-Disposition'] = 'attachment; filename=word_documents.zip'

        # Zip fayl va Word fayllarni o'chirish
        for file_path in word_files:
            os.remove(file_path)
        os.remove(zip_file_path)  # Zip faylni o'chirish

        return response

    return render(request, 'student_excel_word.html')


@login_required
def upload_students(request):
    if request.method == 'POST' and request.FILES['excel_file']:
        excel_file = request.FILES['excel_file']

        # Faylni vaqtinchalik saqlash va load_students_from_excel funksiyasiga uzatish
        load_students_from_excel(excel_file)

        return HttpResponse("Ma'lumotlar bazaga yuklandi!")
    return render(request, 'upload.html')


class HonadonCreateView(LoginRequiredMixin, CreateView):
    model = Honadon
    fields = ['image_1', 'student_name', 'image_2', 'dicription']
    template_name = 'honodon.html'
    success_url = reverse_lazy("honadon-list")  # Muvaffaqiyatli saqlanganidan so'ng

    def form_valid(self, form):
        response = super().form_valid(form)
        self.create_pdf(self.object)  # PDF yaratish
        return response

    def create_pdf(self, honadon):
        pdf = FPDF(orientation='P', unit='mm', format=(210, 150))
        pdf.add_page()

        # Shrift qo'shish
        pdf.set_font("Arial", size=12)

        # Rasm qo'shish
        dir_base = Path(__file__).parent.parent
        image_path = dir_base / "media/photo_2024-10-06_16-45-44.jpg"

        # Rasmning mavjudligini tekshirish
        if os.path.exists(image_path):
            pdf.image(str(image_path), x=0, y=0, w=210, h=150)
        else:
            return HttpResponse("Rasm topilmadi: photo_2024-10-06_16-45-44.jpg", status=404)

        # Birinchi rasmni qo'shish
        pdf.image(honadon.image_1.path, x=5, y=5, w=90, h=60)

        # Birinchi matnni rasmning o'ng tomoniga joylashtirish
        pdf.set_xy(100, 10)
        pdf.multi_cell(100, 10, txt=honadon.student_name, align='L')

        # Matn va rasm orasida bo'sh joy yaratish
        pdf.ln(50)

        # Ikkinchi matnni qo'shish
        pdf.multi_cell(100, 10, txt=honadon.dicription, align='L')

        # Ikkinchi rasmni qo'shish
        pdf.image(honadon.image_2.path, x=110, y=70, w=90, h=60)

        # PDFni saqlash uchun katalog yo'lini tekshirish
        pdf_dir = dir_base / "media/honadon_pdf"
        if not os.path.exists(pdf_dir):
            os.makedirs(pdf_dir)  # Agar katalog mavjud bo'lmasa, yarating

        # PDFni saqlash
        pdf_output_path = pdf_dir / f"honadon_{honadon.pk}.pdf"
        pdf.output(str(pdf_output_path))

        return pdf_output_path  # PDF fayl yo'lini qaytarish


class PDFDownloadView(LoginRequiredMixin, View):  # View dan meros olish
    def get(self, request, pk):
        dir_pdf = Path(__file__).parent.parent
        pdf_file_path = dir_pdf / f"media/honadon_pdf/honadon_{pk}.pdf"  # PDF fayl yo'li

        # PDF faylini yuklab olish
        if os.path.exists(pdf_file_path):
            with open(pdf_file_path, 'rb') as pdf_file:
                response = HttpResponse(pdf_file.read(), content_type='application/pdf')
                response['Content-Disposition'] = f'attachment; filename="honadon_{pk}.pdf"'
                return response
        else:
            return HttpResponse("PDF topilmadi", status=404)


class PDFList(LoginRequiredMixin, ListView):
    model = Honadon
    template_name = 'honadon_list.html'
    context_object_name = 'honadon_list'

    def get_queryset(self):
        return Honadon.objects.all().order_by('-id')  # student_name ga ko'ra kamayish tartibi


class GroupViewCreate(LoginRequiredMixin, UserPassesTestMixin, CreateView):
    model = GroupView
    form_class = GroupViewForm
    template_name = 'group/groupview_form.html'
    success_url = reverse_lazy('groupview_list')

    def test_func(self):
        return self.request.user.is_staff

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        # Shablonga yuborilayotgan barcha ma'lumotlarni konsolga chiqarish
        print(context)  # Konsolga ma'lumotlarni chiqaradi
        return context


class GroupViewUpdate(LoginRequiredMixin, UserPassesTestMixin, UpdateView):
    model = GroupView
    form_class = GroupViewForm
    template_name = 'group/groupview_form.html'  # o'z shabloningizni qo'shing
    success_url = reverse_lazy('groupview_list')  # muvaffaqiyatli yangilashdan so'ng qaytadigan URL

    def test_func(self):
        return self.request.user.is_staff


class GroupViewList(LoginRequiredMixin, UserPassesTestMixin, ListView):
    model = GroupView
    template_name = 'group/grouplist.html'
    context_object_name = 'group_views'

    def test_func(self):
        return self.request.user.is_staff

    def get_queryset(self):
        # Barcha GroupView obyektlarini qaytarish
        return GroupView.objects.all()

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        # Foydalanuvchilar va ularning guruhlarini olish
        user_group_info = []
        for group_view in self.get_queryset():
            user_group_info.append({
                'id': group_view.id,
                'user': group_view.user,
                'groups': group_view.groups_list.all()  # Har bir foydalanuvchiga tegishli guruhlar
            })

        context['user_group_info'] = user_group_info
        return context


